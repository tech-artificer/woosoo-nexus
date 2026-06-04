from pathlib import Path
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
SRC = ROOT / "docs" / "software-development"

DOCUMENTS = [
    ("PROCESS_DOCUMENTATION.md", "woosoo-process-documentation.docx"),
    ("PRODUCT_DOCUMENTATION.md", "woosoo-product-documentation.docx"),
    ("USER_DOCUMENTATION.md", "woosoo-user-documentation.docx"),
]


def strip_frontmatter(text: str) -> str:
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) == 3:
            return parts[2].lstrip()
    return text


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.strip())
    run.bold = bold
    run.font.size = Pt(8.5)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            value = row[c_idx] if c_idx < len(row) else ""
            set_cell_text(table.cell(r_idx, c_idx), value, bold=(r_idx == 0))
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph(style="No Spacing")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
    doc.add_paragraph()


def add_paragraph_with_inline_code(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        run = p.add_run(part[1:-1] if part.startswith("`") and part.endswith("`") else part)
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def build_docx(markdown_path: Path, out_path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    for name, size in [("Heading 1", 18), ("Heading 2", 14), ("Heading 3", 12)]:
        styles[name].font.name = "Arial"
        styles[name].font.size = Pt(size)

    lines = strip_frontmatter(markdown_path.read_text(encoding="utf-8")).splitlines()
    in_code = False
    code_lines: list[str] = []
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            # Drop markdown separator rows.
            cleaned = [
                row for row in table_rows
                if not all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in row)
            ]
            add_table(doc, cleaned)
            table_rows = []

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            flush_table()
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            row = [cell.strip() for cell in line.strip("|").split("|")]
            table_rows.append(row)
            continue
        else:
            flush_table()

        if not line.strip():
            continue

        if line.startswith("# "):
            title = line[2:].strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(20)
            doc.add_paragraph()
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("- [ ]"):
            add_paragraph_with_inline_code(doc, line.replace("- [ ]", "[ ]", 1).strip(), style="List Bullet")
        elif line.startswith("- "):
            add_paragraph_with_inline_code(doc, line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            add_paragraph_with_inline_code(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            add_paragraph_with_inline_code(doc, line)

    flush_table()
    if code_lines:
        add_code_block(doc, code_lines)

    doc.save(out_path)


def verify_outputs() -> None:
    required_phrases = [
        "Tablet sends intent only",
        "order.created",
        "order.updated",
        "order.details.updated",
        "order.printed",
        "session.reset",
        "reserve",
        "ack",
        "failed",
        "pending",
        "confirmed",
        "in_progress",
        "completed",
        "voided",
        "cancelled",
    ]
    forbidden_phrases = [
        "order.print event",
        "print.job event",
        "device_orders.print_event_id column",
    ]

    combined = "\n".join((SRC / md).read_text(encoding="utf-8") for md, _ in DOCUMENTS)
    missing = [phrase for phrase in required_phrases if phrase not in combined]
    forbidden = [phrase for phrase in forbidden_phrases if phrase in combined]
    if missing:
        raise RuntimeError(f"Missing required phrase(s): {missing}")
    if forbidden:
        raise RuntimeError(f"Forbidden stale phrase(s): {forbidden}")

    for _, out_name in DOCUMENTS:
        out = SRC / out_name
        if not out.exists() or out.stat().st_size <= 0:
            raise RuntimeError(f"Missing or empty DOCX: {out}")


def main() -> None:
    for md_name, out_name in DOCUMENTS:
        build_docx(SRC / md_name, SRC / out_name)
        print(f"Wrote {SRC / out_name}")
    verify_outputs()
    print("Structural documentation checks passed.")


if __name__ == "__main__":
    main()
