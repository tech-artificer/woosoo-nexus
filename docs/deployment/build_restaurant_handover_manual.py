from __future__ import annotations

import math
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "deployment" / "woosoo-restaurant-operations-handover-manual.docx"
ASSET_DIR = ROOT / "docs" / "deployment" / "assets"
QA_DIR = ROOT / "docs" / "deployment" / "_rendered_manual_pages"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
PALE = "E8EEF5"
LIGHT = "F4F6F9"
GRID = "9AAEC4"
GREEN = "477A4A"
GOLD = "7A5A00"
RED = "9B1C1C"

COMMAND = "docker compose --env-file ./woosoo-nexus/.env -f compose.yaml"

RESTAURANT_VALUES = {
    "Woosoo server/Pi public host": "192.168.1.31",
    "Krypton Woosoo PC / POS host": "192.168.1.32",
    "Krypton subnet mask": "255.255.255.0",
    "Krypton gateway": "192.168.1.1",
    "POS database": "krypton_woosoo",
    "POS database port": "2121",
}

FORBIDDEN_IPS = [
    "192.168.100.1",
    "192.168.100.7",
    "192.168.100.10",
    "192.168.100.20",
    "192.168.100.42",
]


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if text_size(draw, candidate, fnt)[0] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_box(draw, xy, text, fill, outline, title_font, body_font, radius=18):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=3)
    lines = text.split("\n")
    total = 0
    rendered = []
    for index, line in enumerate(lines):
        fnt = title_font if index == 0 else body_font
        wrapped = wrap_text(draw, line, fnt, x2 - x1 - 36)
        for part in wrapped:
            w, h = text_size(draw, part, fnt)
            rendered.append((part, fnt, w, h))
            total += h + 7
    y = y1 + ((y2 - y1) - total) / 2
    for part, fnt, w, h in rendered:
        draw.text((x1 + (x2 - x1 - w) / 2, y), part, fill="#0B2545", font=fnt)
        y += h + 7


def arrow(draw, start, end, color="#476582", width=4):
    draw.line([start, end], fill=color, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    size = 14
    p1 = (end[0] - size * math.cos(angle - math.pi / 6), end[1] - size * math.sin(angle - math.pi / 6))
    p2 = (end[0] - size * math.cos(angle + math.pi / 6), end[1] - size * math.sin(angle + math.pi / 6))
    draw.polygon([end, p1, p2], fill=color)


def diagram_canvas(title: str):
    image = Image.new("RGB", (1800, 1050), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(48, True)
    draw.text((70, 45), title, fill=f"#{INK}", font=title_font)
    draw.line((70, 115, 1730, 115), fill=f"#{BLUE}", width=5)
    return image, draw


def save_network_diagram():
    image, draw = diagram_canvas("Restaurant Network Topology")
    title_font = font(34, True)
    body_font = font(27)
    small_font = font(23)
    draw_centered_box(draw, (120, 260, 520, 500), "Tablets\nOpen https://192.168.1.31:4443", "#F7FAFF", f"#{GRID}", title_font, body_font)
    draw_centered_box(draw, (700, 190, 1120, 550), "Woosoo Server / Pi\n192.168.1.31\nNexus, Tablet PWA, Reverb, MySQL, Redis", "#EEF5FF", f"#{BLUE}", title_font, body_font)
    draw_centered_box(draw, (1310, 260, 1680, 500), "Krypton Woosoo PC\n192.168.1.32\nMask 255.255.255.0\nGateway 192.168.1.1", "#F8FBF4", f"#{GREEN}", title_font, body_font)
    draw_centered_box(draw, (705, 720, 1115, 900), "Router / Gateway\n192.168.1.1", "#FFF9E8", f"#{GOLD}", title_font, body_font)
    arrow(draw, (520, 380), (700, 380))
    arrow(draw, (1120, 380), (1310, 380))
    arrow(draw, (910, 550), (910, 720))
    draw.text((150, 565), "Customer tablets never send pricing, totals, or POS mapping.", fill=f"#{INK}", font=small_font)
    path = ASSET_DIR / "restaurant-network-topology.png"
    image.save(path)
    return path


def save_responsibility_diagram():
    image, draw = diagram_canvas("Application Responsibility Map")
    title_font = font(31, True)
    body_font = font(24)
    boxes = [
        ((90, 210, 450, 500), "Tablet Ordering PWA\nCustomer flow\nDevice registration\nOrder/refill/service intent"),
        ((530, 210, 890, 500), "Woosoo Nexus\nBusiness truth\nSessions and devices\nOrders and POS writes\nRealtime and print events"),
        ((970, 210, 1330, 500), "Krypton POS\nPOS database\nCurrent session\nMenu/package records\nRestaurant transaction record"),
        ((1410, 210, 1740, 500), "Print Bridge\nPrinter status\nEvent intake\nReserve / ACK / failed"),
    ]
    for xy, text in boxes:
        draw_centered_box(draw, xy, text, "#F6F9FC", f"#{GRID}", title_font, body_font)
    for start, end in [((450, 355), (530, 355)), ((890, 355), (970, 355)), ((1330, 355), (1410, 355))]:
        arrow(draw, start, end)
    draw_centered_box(draw, (300, 700, 1500, 890), "Operator Rule\nBackend owns truth. Tablet sends intent only. Print Bridge proves last-mile print outcome.", "#EEF5FF", f"#{BLUE}", title_font, body_font)
    path = ASSET_DIR / "app-responsibility-map.png"
    image.save(path)
    return path


def save_order_flow_diagram():
    image, draw = diagram_canvas("Order Flow")
    title_font = font(31, True)
    body_font = font(24)
    steps = [
        ((80, 270, 350, 500), "1. Tablet\nGuest count\nPackage\nItems"),
        ((430, 270, 700, 500), "2. Nexus API\nValidate intent\nCompute truth"),
        ((780, 270, 1050, 500), "3. Krypton POS\nWrite order\nUse active session"),
        ((1130, 270, 1400, 500), "4. Print Event\nStation routing\nQueue event"),
        ((1480, 270, 1720, 500), "5. Bridge\nPrint\nACK / failed"),
    ]
    for xy, text in steps:
        draw_centered_box(draw, xy, text, "#F8FAFD", f"#{GRID}", title_font, body_font)
    for start, end in [((350, 385), (430, 385)), ((700, 385), (780, 385)), ((1050, 385), (1130, 385)), ((1400, 385), (1480, 385))]:
        arrow(draw, start, end)
    draw_centered_box(draw, (210, 710, 1590, 890), "Failure Policy\nDo not fabricate success. Customer screens must show safe messages. Technical errors belong in logs.", "#FFF4F4", f"#{RED}", title_font, body_font)
    path = ASSET_DIR / "order-flow.png"
    image.save(path)
    return path


def save_deployment_workflow_diagram():
    image, draw = diagram_canvas("Deployment And Redeployment Workflow")
    title_font = font(29, True)
    body_font = font(23)
    steps = [
        ((100, 220, 430, 420), "Preflight\nsudo bash\nscripts/deployment/\ndoctor.sh"),
        ((540, 220, 870, 420), "Apply Config\nsudo bash\nscripts/deployment/\napply-woosoo-config.sh"),
        ((980, 220, 1310, 420), "Deploy\nsudo bash\nscripts/deployment/\ndeploy.sh"),
        ((1420, 220, 1720, 420), "Verify\nps, logs, smoke checks"),
    ]
    for xy, text in steps:
        draw_centered_box(draw, xy, text, "#F7FAFF", f"#{BLUE}", title_font, body_font)
    for start, end in [((430, 320), (540, 320)), ((870, 320), (980, 320)), ((1310, 320), (1420, 320))]:
        arrow(draw, start, end)
    draw_centered_box(draw, (260, 690, 1540, 900), "Rollback Rule\nNever use docker compose down -v for normal rollback. Preserve data volumes and use known-good refs/backups.", "#FFF9E8", f"#{GOLD}", title_font, body_font)
    path = ASSET_DIR / "deployment-workflow.png"
    image.save(path)
    return path


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2EC"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, bold=False, size=9.5, color="0B2545"):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, PALE)
        set_cell_border(cell)
        set_cell_text(cell, header, bold=True, size=9.5, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_border(cells[i])
            set_cell_text(cells[i], str(value), size=9.2)
    doc.add_paragraph()
    return table


def add_code(doc, text):
    p = doc.add_paragraph(style="CodeBlock")
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    run.font.size = Pt(8.4)
    return p


def add_note(doc, title, body, kind="info"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    fill = {"info": LIGHT, "warn": "FFF2CC", "risk": "FDEAEA"}.get(kind, LIGHT)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "C9D6E2")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    title_run = p.add_run(title + "\n")
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(INK)
    body_run = p.add_run(body)
    body_run.font.size = Pt(9.3)
    doc.add_paragraph()


def add_placeholder(doc, title, proof):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F8FAFC")
    set_cell_border(cell, "7A8FA6")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SCREENSHOT PLACEHOLDER\n")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    r2 = p.add_run(title + "\n")
    r2.bold = True
    r3 = p.add_run("Proof required: " + proof)
    r3.font.size = Pt(9)
    doc.add_paragraph()


def set_doc_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 14, 7),
        ("Heading 2", 13, BLUE, 11, 5),
        ("Heading 3", 11.5, DARK_BLUE, 8, 3),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(8.4)
    code.paragraph_format.left_indent = Inches(0.15)
    code.paragraph_format.right_indent = Inches(0.15)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)
    code.paragraph_format.line_spacing = 1.05


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)


def add_header_footer(doc: Document):
    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = "Woosoo Restaurant Operations And Handover Manual"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if header.runs:
            header.runs[0].font.size = Pt(8.5)
            header.runs[0].font.color.rgb = RGBColor.from_string("6B7280")
        footer = section.footer.paragraphs[0]
        footer.text = "Restaurant network only - no test network values"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if footer.runs:
            footer.runs[0].font.size = Pt(8.5)
            footer.runs[0].font.color.rgb = RGBColor.from_string("6B7280")


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Woosoo Restaurant Operations And Handover Manual")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(INK)
    subtitle = doc.add_paragraph()
    subtitle.add_run("Deployment, operations, troubleshooting, rollback, and acceptance guide").italic = True
    meta = doc.add_paragraph()
    meta.add_run("Prepared: 2026-05-22\nScope: Restaurant production LAN\nDeliverable: DOCX handover manual with screenshot placeholders")
    add_note(
        doc,
        "Operator Safety",
        "This manual is documentation-only. It does not change app code, Docker behavior, API contracts, or order logic. Never expose secrets in screenshots or chat.",
        "warn",
    )
    add_table(
        doc,
        ["Network Item", "Restaurant Value"],
        list(RESTAURANT_VALUES.items()),
        [2.7, 3.8],
    )
    p = doc.add_paragraph()
    p.add_run("Source of truth: ").bold = True
    p.add_run("Use this restaurant manual and the platform-root deployment docs for handover. Do not copy non-restaurant values from older examples.")
    doc.add_page_break()


def add_business_requirements(doc):
    add_heading(doc, "1. Business Requirements", 1)
    for item in [
        "Restaurant ordering must run on the local LAN during normal service.",
        "Backend owns truth: Woosoo Nexus owns pricing, packages, sessions, orders, POS/Krypton integration, realtime events, and print events.",
        "The Tablet Ordering PWA sends customer intent only and must not send pricing, tax, totals, POS mapping, or backend state.",
        "The Print Bridge proves last-mile printer outcome through heartbeat, reserve, ACK, and failed flows.",
        "Operators need repeatable deployment, log checking, troubleshooting, rollback, and handover acceptance steps.",
        "Customer-facing screens must never show raw technical errors.",
    ]:
        bullet(doc, item)


def add_specs(doc):
    add_heading(doc, "2. Restaurant Network And System Specifications", 1)
    add_table(doc, ["Specification", "Value"], list(RESTAURANT_VALUES.items()), [2.5, 4.0])
    add_note(
        doc,
        "POS Port Lock",
        "Restaurant deployment uses POS database port 2121. Confirm this value in the live restaurant config before cutover and do not substitute older sample ports.",
        "warn",
    )
    add_table(
        doc,
        ["Service", "Published Endpoint", "Purpose"],
        [
            ("Nexus admin/API", "https://192.168.1.31", "Admin UI and backend API"),
            ("Tablet PWA", "https://192.168.1.31:4443", "Customer tablet app"),
            ("Reverb WebSocket", "wss://192.168.1.31/app/{REVERB_APP_KEY}", "Realtime events through nginx"),
            ("Krypton POS DB", "192.168.1.32:2121", "External POS database connection"),
        ],
        [1.7, 2.4, 2.4],
    )


def add_diagrams(doc, paths):
    add_heading(doc, "3. Generated Diagrams", 1)
    captions = [
        ("Restaurant network topology", paths["network"]),
        ("Application responsibility map", paths["responsibility"]),
        ("Order flow", paths["order"]),
        ("Deployment and redeployment workflow", paths["deployment"]),
    ]
    for caption, path in captions:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
        doc.add_picture(str(path), width=Inches(6.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()


def add_app_manuals(doc):
    add_heading(doc, "4. App Manuals", 1)
    manuals = [
        (
            "Woosoo Nexus",
            [
                "Purpose: Laravel admin panel, backend API, POS/Krypton integration, sessions, orders, Reverb, and print events.",
                "Operators use it for device setup, order monitoring, menu/package management, and operational checks.",
                "Check logs first in app, nginx, reverb, mysql, and redis before changing configuration.",
            ],
        ),
        (
            "Tablet Ordering PWA",
            [
                "Purpose: customer-facing tablet ordering flow.",
                "Flow: welcome, device registration/settings, guest count, package selection, menu/cart, review, submit, in-session refill/service request.",
                "Tablet sends intent only; backend computes truth.",
                "Installed tablets update after deployment through service-worker/build-info checks; staff may force update from settings when needed.",
            ],
        ),
        (
            "Woosoo Print Bridge",
            [
                "Purpose: Android relay for actual printer dispatch.",
                "The bridge reports heartbeat and acknowledges print outcome.",
                "Print lifecycle is reserve, ack, or failed; retry must not create duplicate physical prints.",
            ],
        ),
        (
            "Krypton Woosoo PC / POS",
            [
                "Purpose: external POS database and restaurant transaction record.",
                "Static IPv4 must be set to 192.168.1.32 with mask 255.255.255.0 and gateway 192.168.1.1.",
                "Connectivity must be verified before deployment acceptance.",
            ],
        ),
        (
            "Platform Docker Stack",
            [
                "Purpose: runtime orchestration from /opt/woosoo/woosoo-platform.",
                "compose.yaml is the runtime authority.",
                "Use the platform-root compose command with --env-file ./woosoo-nexus/.env.",
            ],
        ),
    ]
    for title, items in manuals:
        add_heading(doc, title, 2)
        for item in items:
            bullet(doc, item)


def add_navigation_guide(doc):
    add_heading(doc, "5. How To Navigate The Apps", 1)
    add_heading(doc, "Woosoo Nexus Admin", 2)
    doc.add_paragraph(
        "After login, use the left sidebar. Admin users see Main, Analytics, and Configuration groups."
    )
    add_table(
        doc,
        ["Group", "Pages", "When to use"],
        [
            (
                "Main",
                "Dashboard, Orders, POS, Menus, Packages, User Management, Devices, Service Requests",
                "Daily operations, device setup, order monitoring, staff/user work, menu/package updates",
            ),
            (
                "Analytics",
                "Reports, Daily Sales, Hourly Sales, Guest Count, Menu Items, Order Status, Print Audit, Discount & Tax",
                "End-of-day checks, sales analysis, print audit, guest and menu reporting",
            ),
            (
                "Configuration",
                "Branches, Access Control, Accessibility, Event Logs, Reverb Service, Monitoring",
                "System setup, role/permission control, audit logs, realtime health, queue/database checks",
            ),
        ],
        [1.25, 2.85, 2.4],
    )
    for item in [
        "Use Dashboard for the live overview.",
        "Use Orders for live orders and order history.",
        "Use Devices for tablet and relay setup.",
        "Use Monitoring for queue, database, and Reverb health.",
        "Use Manual for the in-app guide library.",
        "Use POS for POS-specific terminal, table, and order inspection.",
    ]:
        bullet(doc, item)

    add_heading(doc, "Tablet Ordering PWA", 2)
    doc.add_paragraph("The customer-facing tablet path is:")
    for step in [
        "Open the welcome screen at /.",
        "If the tablet is not registered, open Settings from the gear icon, enter or create the PIN, and complete registration.",
        "Tap Begin the Feast.",
        "Select guest count on /order/start.",
        "Choose a dining package on /order/packageSelection.",
        "Browse meats, sides, desserts, and drinks on /menu.",
        "Open the order summary and continue to /order/review.",
        "Submit the order; the app sends POST /api/devices/create-order.",
        "Use /order/in-session for submitted items, refills, service requests, and session status.",
        "When the session ends, /order/session-ended returns the tablet toward the welcome screen for the next table.",
    ]:
        numbered(doc, step)
    add_note(
        doc,
        "Staff-only tablet maintenance",
        "Tablet settings live under /settings. The /sw-reset route is an emergency same-origin reset path and should be used only for dedicated-origin recovery when normal settings maintenance cannot refresh the active tablet app.",
        "warn",
    )


def add_directories_and_commands(doc):
    add_heading(doc, "6. Directory Structure And Common Commands", 1)
    add_code(
        doc,
        "/opt/woosoo/woosoo-platform/\n"
        "  compose.yaml\n"
        "  docker/\n"
        "    certs/\n"
        "    nginx/\n"
        "    mysql/\n"
        "    php/\n"
        "  scripts/deployment/\n"
        "  woosoo-nexus/\n"
        "  tablet-ordering-pwa/\n"
        "/etc/woosoo/woosoo.env",
    )
    add_heading(doc, "Navigation", 2)
    add_code(doc, "cd /opt/woosoo/woosoo-platform\npwd\nls -la\ncd woosoo-nexus\ncd ../tablet-ordering-pwa\ncd ..")
    add_heading(doc, "Docker Status And Logs", 2)
    add_code(
        doc,
        f"{COMMAND} ps\n"
        f"{COMMAND} logs --tail=100 app\n"
        f"{COMMAND} logs --tail=100 nginx\n"
        f"{COMMAND} logs --tail=100 reverb\n"
        f"{COMMAND} logs --tail=100 mysql redis",
    )
    add_heading(doc, "Deployment Commands", 2)
    add_code(doc, "sudo bash scripts/deployment/doctor.sh\nsudo bash scripts/deployment/apply-woosoo-config.sh\nsudo bash scripts/deployment/deploy.sh")
    add_heading(doc, "Smoke Checks", 2)
    add_code(
        doc,
        "ping 192.168.1.32\n"
        "curl -k https://192.168.1.31\n"
        "curl -k https://192.168.1.31:4443/build-info.json\n"
        f"{COMMAND} exec -T app php artisan route:list\n"
        f"{COMMAND} exec -T app php artisan config:clear",
    )


def add_workflows(doc):
    add_heading(doc, "7. Operational Workflows", 1)
    flows = [
        (
            "First-Time Restaurant Setup",
            [
                "Set Krypton Woosoo PC IPv4 to the required restaurant values.",
                "Confirm Woosoo server/Pi is reachable at 192.168.1.31.",
                "Confirm /etc/woosoo/woosoo.env exists and contains restaurant production values with secrets hidden from screenshots.",
                "Run deployment doctor and resolve all blocking errors.",
                "Apply config, deploy, then verify Nexus, Tablet PWA, POS, Reverb, and print flow.",
            ],
        ),
        (
            "Daily Startup Check",
            [
                "Run compose ps and confirm expected services are running.",
                "Open Nexus admin and Tablet PWA.",
                "Ping the POS host.",
                "Review app, nginx, reverb, mysql, and redis logs for blockers.",
                "Confirm Print Bridge is online before service starts.",
            ],
        ),
        (
            "Redeployment",
            [
                "Run deployment doctor.",
                "Confirm no uncommitted emergency changes are present on deployed repos.",
                "Run deploy.sh from platform root.",
                "Check service health and build-info.",
                "Run final acceptance checks.",
            ],
        ),
        (
            "Rollback",
            [
                "Preserve Docker volumes. Do not use down -v for normal rollback.",
                "Return app repos to the previous known-good refs.",
                "Rebuild and start only the affected services when possible.",
                "Verify Nexus, Tablet PWA, POS connectivity, Reverb, and print flow again.",
            ],
        ),
    ]
    for title, steps in flows:
        add_heading(doc, title, 2)
        for step in steps:
            numbered(doc, step)


def add_screenshot_placeholders(doc):
    add_heading(doc, "8. Screenshot Placeholders", 1)
    placeholders = [
        ("Windows adapter list", "Active Ethernet adapter is selected on the Krypton Woosoo PC."),
        ("Krypton PC IPv4 properties", "IPv4 shows IP 192.168.1.32, mask 255.255.255.0, gateway 192.168.1.1."),
        ("Restaurant environment config", "/etc/woosoo/woosoo.env shows restaurant values and hides secrets."),
        ("Docker service status", "docker compose ps shows expected services running or healthy."),
        ("Nexus admin", "Nexus opens at https://192.168.1.31."),
        ("Tablet PWA", "Tablet app opens at https://192.168.1.31:4443."),
        ("POS connectivity", "Ping or equivalent check reaches 192.168.1.32."),
        ("Print Bridge", "Bridge status screen shows printer/relay online before service."),
    ]
    for title, proof in placeholders:
        add_placeholder(doc, title, proof)


def add_troubleshooting(doc):
    add_heading(doc, "9. Troubleshooting", 1)
    add_table(
        doc,
        ["Symptom", "First Check", "Command / Action"],
        [
            ("Nexus admin unreachable", "nginx and app logs", f"{COMMAND} logs --tail=100 nginx app"),
            ("Tablet app unreachable", "tablet endpoint", "curl -k https://192.168.1.31:4443/build-info.json"),
            ("POS connection fails", "Krypton PC network", "ping 192.168.1.32"),
            ("Realtime/WebSocket fails", "reverb and app logs", f"{COMMAND} logs --tail=100 reverb app"),
            ("MySQL/Redis unhealthy", "service logs", f"{COMMAND} logs --tail=100 mysql redis"),
            ("Tablet stuck on old build", "build-info", "curl -k https://192.168.1.31:4443/build-info.json"),
            ("Orders do not print", "Bridge status and print events", "Check bridge app and Nexus print-event logs"),
            ("Deployment script fails", "preflight config", "sudo bash scripts/deployment/doctor.sh"),
        ],
        [1.7, 1.7, 3.1],
    )


def add_handover(doc):
    add_heading(doc, "10. Final Handover Acceptance", 1)
    for item in [
        "Nexus reachable at https://192.168.1.31.",
        "Tablet PWA reachable at https://192.168.1.31:4443.",
        "Krypton POS host 192.168.1.32 reachable.",
        "A test order reaches POS.",
        "Print flow verified by bridge ACK or operator-confirmed print.",
        "Logs reviewed with no unresolved deployment blockers.",
        "Operator knows where /etc/woosoo/woosoo.env, compose.yaml, docker/certs, and scripts/deployment live.",
        "Rollback path and backup location explained.",
        "Screenshots inserted or queued for insertion without exposing secrets.",
    ]:
        bullet(doc, "[ ] " + item)


def verify_docx_text(path: Path):
    with zipfile.ZipFile(path) as zf:
        text = "\n".join(
            zf.read(name).decode("utf-8", errors="ignore")
            for name in zf.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )
    missing = [value for value in RESTAURANT_VALUES.values() if value not in text]
    forbidden = [value for value in FORBIDDEN_IPS if value in text]
    required_phrases = [
        "SCREENSHOT PLACEHOLDER",
        "docker compose --env-file ./woosoo-nexus/.env -f compose.yaml ps",
        "Final Handover Acceptance",
        "Backend owns truth",
        "Begin the Feast",
        "Woosoo Nexus Admin",
    ]
    missing_phrases = [phrase for phrase in required_phrases if phrase not in text]
    if missing or forbidden or missing_phrases:
        raise RuntimeError(f"manual verification failed: missing={missing}, forbidden={forbidden}, missing_phrases={missing_phrases}")


def build():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    if QA_DIR.exists():
        shutil.rmtree(QA_DIR)
    paths = {
        "network": save_network_diagram(),
        "responsibility": save_responsibility_diagram(),
        "order": save_order_flow_diagram(),
        "deployment": save_deployment_workflow_diagram(),
    }

    doc = Document()
    set_doc_styles(doc)
    add_cover(doc)
    add_business_requirements(doc)
    add_specs(doc)
    add_diagrams(doc, paths)
    add_app_manuals(doc)
    add_navigation_guide(doc)
    add_directories_and_commands(doc)
    add_workflows(doc)
    add_screenshot_placeholders(doc)
    add_troubleshooting(doc)
    add_handover(doc)
    add_header_footer(doc)
    doc.save(OUT)
    verify_docx_text(OUT)
    print(f"Wrote {OUT}")
    for name, path in paths.items():
        print(f"Wrote {name}: {path}")


if __name__ == "__main__":
    build()
